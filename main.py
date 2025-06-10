# AI-Based Job Screening System
# Main implementation with core components

import os
import re
import json
import numpy as np
import pandas as pd
import spacy
import PyPDF2
import docx
import pdfplumber
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Load NLP model
try:
    nlp = spacy.load("en_core_web_sm")  # Changed to smaller model that's more commonly available
except OSError:
    print("Please install spacy English model: python -m spacy download en_core_web_sm")
    nlp = None

class JobDescriptionParser:
    """Extracts key requirements from job descriptions"""
    
    def __init__(self):  # Fixed: __init__ instead of _init_
        # Removed T5 model dependency for now
        pass
    
    def extract_text(self, file_path):
        """Extract text from JD files in various formats"""
        if file_path.endswith('.pdf'):
            return self._extract_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            return self._extract_from_docx(file_path)
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        else:
            raise ValueError("Unsupported file format")
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"Error extracting from PDF: {e}")
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e2:
                print(f"Fallback PDF extraction also failed: {e2}")
        return text
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            print(f"Error extracting from DOCX: {e}")
            return ""
    
    def summarize(self, text):
        """Generate simple summary of job description"""
        # Simple summarization without T5 model
        sentences = text.split('.')
        # Take first few sentences as summary
        summary = '. '.join(sentences[:3]) + '.'
        return summary
    
    def extract_requirements(self, text):
        """Extract key requirements from job description"""
        if not nlp:
            return self._extract_requirements_basic(text)
        
        doc = nlp(text)
        
        # Extract skills, education, experience
        skills = []
        education = []
        experience = []
        
        # Basic pattern matching for common requirement indicators
        skill_patterns = ["skills", "proficiency", "knowledge of", "familiarity with", "expertise in"]
        edu_patterns = ["degree", "qualification", "education", "graduate", "bachelor", "master", "phd"]
        exp_patterns = ["experience", "years", "worked", "background in"]
        
        for sent in doc.sents:
            sent_text = sent.text.lower()
            
            # Check for skills
            if any(pattern in sent_text for pattern in skill_patterns):
                skills.append(sent.text.strip())
            
            # Check for education
            if any(pattern in sent_text for pattern in edu_patterns):
                education.append(sent.text.strip())
            
            # Check for experience
            if any(pattern in sent_text for pattern in exp_patterns):
                experience.append(sent.text.strip())
        
        # Extract job location and type from the text
        location_pattern = r"location:?\s*([\w\s,]+)"
        job_type_pattern = r"(full[- ]time|part[- ]time|contract|remote|hybrid)"
        
        location_match = re.search(location_pattern, text, re.IGNORECASE)
        job_type_match = re.search(job_type_pattern, text, re.IGNORECASE)
        
        location = location_match.group(1).strip() if location_match else "Not specified"
        job_type = job_type_match.group(1).strip() if job_type_match else "Not specified"
        
        # Organize extracted information
        requirements = {
            "skills": skills,
            "education": education,
            "experience": experience,
            "location": location,
            "job_type": job_type,
            "summary": self.summarize(text)
        }
        
        return requirements
    
    def _extract_requirements_basic(self, text):
        """Basic requirements extraction without spaCy"""
        text_lower = text.lower()
        
        # Extract skills using common patterns
        skills = []
        skill_keywords = ["python", "java", "javascript", "sql", "html", "css", "react", "angular", "node"]
        for skill in skill_keywords:
            if skill in text_lower:
                skills.append(skill)
        
        # Basic education extraction
        education = []
        if "bachelor" in text_lower or "degree" in text_lower:
            education.append("Bachelor's degree required")
        if "master" in text_lower:
            education.append("Master's degree preferred")
        
        # Basic experience extraction
        experience = []
        exp_match = re.search(r"(\d+)\s*years?\s*experience", text_lower)
        if exp_match:
            experience.append(f"{exp_match.group(1)} years of experience required")
        
        return {
            "skills": skills,
            "education": education,
            "experience": experience,
            "location": "Not specified",
            "job_type": "Not specified",
            "summary": self.summarize(text)
        }


class ResumeParser:
    """Extracts structured information from candidate resumes"""
    
    def __init__(self):  # Fixed: __init__ instead of _init_
        self.nlp = nlp
    
    def extract_text(self, file_path):
        """Extract text from resume in various formats"""
        if file_path.endswith('.pdf'):
            return self._extract_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            return self._extract_from_docx(file_path)
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        else:
            raise ValueError("Unsupported file format")
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"Error extracting from PDF: {e}")
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e2:
                print(f"Fallback PDF extraction also failed: {e2}")
        return text
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            print(f"Error extracting from DOCX: {e}")
            return ""
    
    def extract_contact_info(self, text):
        """Extract contact information from resume"""
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        
        # Phone pattern
        phone_pattern = r'(\+\d{1,3}\s?)?(\(?\d{3}\)?[\s.-]?)?\d{3}[\s.-]?\d{4}'
        phones = re.findall(phone_pattern, text)
        
        # Name extraction (basic approach)
        name = ""
        if self.nlp:
            doc = self.nlp(text)
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    name = ent.text
                    break
        
        # If no name found, try to get from first line
        if not name:
            lines = text.strip().split('\n')
            if lines:
                first_line = lines[0].strip()
                if len(first_line.split()) <= 3 and not any(char.isdigit() for char in first_line):
                    name = first_line
        
        return {
            "name": name,
            "email": emails[0] if emails else "",
            "phone": ''.join(phones[0]) if phones else ""
        }
    
    def extract_education(self, text):
        """Extract education information from resume"""
        education = []
        edu_keywords = ["degree", "bachelor", "master", "phd", "btech", "mtech", "b.tech", "m.tech", "diploma"]
        
        if self.nlp:
            doc = self.nlp(text)
            for sent in doc.sents:
                sent_text = sent.text.lower()
                if any(keyword in sent_text for keyword in edu_keywords):
                    education.append(sent.text.strip())
        else:
            # Basic approach without spaCy
            lines = text.split('\n')
            for line in lines:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in edu_keywords):
                    education.append(line.strip())
        
        return education
    
    def extract_skills(self, text):
        """Extract skills from resume"""
        # Common technical skills
        tech_skills = [
            "python", "java", "javascript", "sql", "html", "css", "react", "angular", "node", 
            "aws", "azure", "gcp", "docker", "kubernetes", "git", "tensorflow", "pytorch", 
            "machine learning", "data science", "deep learning", "ai", "artificial intelligence",
            "nlp", "natural language processing", "computer vision", "c++", "c#", "php", "r"
        ]
        
        skills = []
        text_lower = text.lower()
        
        for skill in tech_skills:
            if skill in text_lower:
                skills.append(skill)
        
        return skills
    
    def extract_experience(self, text):
        """Extract work experience from resume"""
        experience = []
        exp_keywords = ["experience", "work", "employment", "job", "position", "role"]
        
        if self.nlp:
            doc = self.nlp(text)
            
            # Try to find experience section
            exp_section = None
            lines = text.split('\n')
            for i, line in enumerate(lines):
                if any(keyword in line.lower() for keyword in exp_keywords) and len(line) < 30:
                    exp_section = i
                    break
            
            if exp_section is not None:
                # Extract text from experience section
                exp_text = '\n'.join(lines[exp_section+1:])
                # Extract sentences with company names or job titles
                for sent in self.nlp(exp_text).sents:
                    for ent in sent.ents:
                        if ent.label_ == "ORG":  # Organization entity
                            experience.append(sent.text.strip())
                            break
        else:
            # Basic approach without spaCy
            lines = text.split('\n')
            for line in lines:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in exp_keywords):
                    experience.append(line.strip())
        
        return experience
    
    def parse_resume(self, file_path):
        """Parse a resume file and extract structured information"""
        text = self.extract_text(file_path)
        
        contact_info = self.extract_contact_info(text)
        education = self.extract_education(text)
        skills = self.extract_skills(text)
        experience = self.extract_experience(text)
        
        return {
            "contact_info": contact_info,
            "education": education,
            "skills": skills,
            "experience": experience,
            "full_text": text
        }


class MatchingEngine:
    """Implements algorithms for matching resumes against job requirements"""
    
    def __init__(self):  # Fixed: __init__ instead of _init_
        self.vectorizer = TfidfVectorizer(ngram_range=(1, 2))
    
    def calculate_skill_match(self, jd_skills, resume_skills):
        """Calculate match score for skills"""
        if not jd_skills:
            return 1.0  # No skills required means 100% match
        
        # Extract skill keywords from JD skills text
        jd_skill_keywords = []
        for skill_text in jd_skills:
            words = skill_text.lower().split()
            jd_skill_keywords.extend(words)
        
        matches = 0
        for skill in jd_skill_keywords:
            if any(skill in resume_skill.lower() for resume_skill in resume_skills):
                matches += 1
        
        return matches / len(jd_skill_keywords) if jd_skill_keywords else 0.0
    
    def calculate_experience_match(self, jd_experience, resume_experience):
        """Calculate match score for experience using semantic similarity"""
        if not jd_experience or not resume_experience:
            return 0.0
            
        # Combine all experience text
        jd_exp_text = ' '.join(jd_experience)
        resume_exp_text = ' '.join(resume_experience)
        
        try:
            # Create TF-IDF vectors
            corpus = [jd_exp_text, resume_exp_text]
            tfidf_matrix = self.vectorizer.fit_transform(corpus)
            
            # Calculate cosine similarity
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            return similarity
        except Exception as e:
            print(f"Error calculating experience match: {e}")
            return 0.0
    
    def calculate_education_match(self, jd_education, resume_education):
        """Calculate match score for education"""
        if not jd_education:
            return 1.0  # No education requirements means 100% match
            
        if not resume_education:
            return 0.0
        
        try:
            # Combined education text
            jd_edu_text = ' '.join(jd_education)
            resume_edu_text = ' '.join(resume_education)
            
            # Create TF-IDF vectors
            corpus = [jd_edu_text, resume_edu_text]
            tfidf_matrix = self.vectorizer.fit_transform(corpus)
            
            # Calculate cosine similarity
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            return similarity
        except Exception as e:
            print(f"Error calculating education match: {e}")
            return 0.0
    
    def match_resume_to_jd(self, job_requirements, resume_data, weights=None):
        """Calculate overall match score between resume and job description"""
        # Default weights if not provided
        if weights is None:
            weights = {
                "skills": 0.5,
                "experience": 0.3,
                "education": 0.2
            }
        
        # Calculate individual scores
        skill_score = self.calculate_skill_match(
            job_requirements.get("skills", []),
            resume_data.get("skills", [])
        )
        
        experience_score = self.calculate_experience_match(
            job_requirements.get("experience", []),
            resume_data.get("experience", [])
        )
        
        education_score = self.calculate_education_match(
            job_requirements.get("education", []),
            resume_data.get("education", [])
        )
        
        # Calculate weighted score
        overall_score = (
            skill_score * weights["skills"] +
            experience_score * weights["experience"] +
            education_score * weights["education"]
        )
        
        # Detailed breakdown
        match_details = {
            "overall_score": overall_score,
            "skill_score": skill_score,
            "experience_score": experience_score,
            "education_score": education_score,
            "matched_skills": []
        }
        
        # Find matched skills
        jd_skills = job_requirements.get("skills", [])
        resume_skills = resume_data.get("skills", [])
        
        for jd_skill in jd_skills:
            jd_skill_words = jd_skill.lower().split()
            for word in jd_skill_words:
                if any(word in resume_skill.lower() for resume_skill in resume_skills):
                    match_details["matched_skills"].append(word)
        
        return match_details


class NotificationSystem:
    """Handles automated communications with candidates"""
    
    def __init__(self, smtp_server, smtp_port, sender_email, sender_password):  # Fixed: __init__
        self.smtp_server = smtp_server
        self.smtp_port = smtp_port
        self.sender_email = sender_email
        self.sender_password = sender_password
        
        # Email templates
        self.templates = {
            "shortlisted": """
            Subject: You've Been Shortlisted for {job_title}
            
            Dear {candidate_name},
            
            We're pleased to inform you that your application for the {job_title} position 
            has been shortlisted. Our AI-powered screening system identified a strong match 
            between your qualifications and our requirements.
            
            Your overall match score: {match_score}%
            
            Our HR team will contact you shortly to schedule an interview.
            
            Best regards,
            {company_name} Recruitment Team
            """,
            
            "rejected": """
            Subject: Update on Your Application for {job_title}
            
            Dear {candidate_name},
            
            Thank you for your interest in the {job_title} position at {company_name}.
            
            After careful review of your application, we've decided to proceed with other 
            candidates whose qualifications more closely match our current requirements.
            
            We encourage you to apply for future openings that align with your skills.
            
            Best regards,
            {company_name} Recruitment Team
            """
        }
    
    def send_email(self, recipient_email, template_name, params):
        """Send email using selected template and parameters"""
        template = self.templates.get(template_name)
        if not template:
            raise ValueError(f"Template '{template_name}' not found")
        
        # Extract subject line and format email content
        lines = template.strip().split('\n')
        subject_line = lines[1].replace('Subject: ', '').strip().format(**params)
        body = '\n'.join(lines[3:]).format(**params)
        
        # Create message
        message = MIMEMultipart()
        message["From"] = self.sender_email
        message["To"] = recipient_email
        message["Subject"] = subject_line
        message.attach(MIMEText(body, "plain"))
        
        try:
            # Connect to SMTP server and send email
            with smtplib.SMTP(self.smtp_server, self.smtp_port) as server:
                server.starttls()
                server.login(self.sender_email, self.sender_password)
                server.send_message(message)
            return True
        except Exception as e:
            print(f"Failed to send email: {e}")
            return False
    
    def notify_candidate(self, candidate_info, job_info, match_details, threshold=0.8):
        """Notify candidate based on match score"""
        match_score = match_details["overall_score"] * 100
        recipient_email = candidate_info["contact_info"]["email"]
        candidate_name = candidate_info["contact_info"]["name"] or "Candidate"
        
        params = {
            "candidate_name": candidate_name,
            "job_title": job_info.get("title", "the position"),
            "company_name": job_info.get("company", "our company"),
            "match_score": round(match_score, 1)
        }
        
        if match_score >= (threshold * 100):
            return self.send_email(recipient_email, "shortlisted", params)
        else:
            return self.send_email(recipient_email, "rejected", params)


class JobScreeningSystem:
    """Main system class that coordinates the job screening process"""
    
    def __init__(self, smtp_config=None):  # Fixed: __init__ instead of _init_
        self.jd_parser = JobDescriptionParser()
        self.resume_parser = ResumeParser()
        self.matching_engine = MatchingEngine()
        
        if smtp_config:
            self.notification_system = NotificationSystem(**smtp_config)
        else:
            self.notification_system = None
        
        self.job_requirements = None
        self.candidates = []
        self.match_results = []
        
    def load_job_description(self, file_path):
        """Load and parse job description"""
        text = self.jd_parser.extract_text(file_path)
        self.job_requirements = self.jd_parser.extract_requirements(text)
        
        # Add job title from filename or first line
        filename = os.path.basename(file_path)
        job_title = os.path.splitext(filename)[0]
        self.job_requirements["title"] = job_title
        
        return self.job_requirements
    
    def load_resumes(self, directory_path):
        """Load and parse all resumes in directory"""
        self.candidates = []
        
        if not os.path.exists(directory_path):
            print(f"Directory {directory_path} does not exist")
            return []
        
        for filename in os.listdir(directory_path):
            if filename.endswith(('.pdf', '.docx', '.txt')):
                file_path = os.path.join(directory_path, filename)
                try:
                    resume_data = self.resume_parser.parse_resume(file_path)
                    resume_data["file_name"] = filename
                    self.candidates.append(resume_data)
                    print(f"Successfully parsed: {filename}")
                except Exception as e:
                    print(f"Error parsing resume {filename}: {e}")
        
        return self.candidates
    
    def screen_candidates(self, threshold=0.8, weights=None):
        """Screen all candidates against job requirements"""
        if not self.job_requirements:
            raise ValueError("Job requirements not loaded")
        
        if not self.candidates:
            raise ValueError("No candidates loaded")
        
        self.match_results = []
        
        for candidate in self.candidates:
            match_details = self.matching_engine.match_resume_to_jd(
                self.job_requirements,
                candidate,
                weights
            )
            
            result = {
                "candidate": candidate["contact_info"]["name"] or "Unknown",
                "email": candidate["contact_info"]["email"],
                "file_name": candidate["file_name"],
                "match_score": match_details["overall_score"],
                "shortlisted": match_details["overall_score"] >= threshold,
                "details": match_details
            }
            
            self.match_results.append(result)
        
        # Sort results by match score in descending order
        self.match_results.sort(key=lambda x: x["match_score"], reverse=True)
        
        return self.match_results
    
    def send_notifications(self, company_name=None):
        """Send notifications to all candidates"""
        if not self.notification_system:
            raise ValueError("Notification system not configured")
        
        if not self.match_results:
            raise ValueError("No screening results available")
        
        job_info = {
            "title": self.job_requirements.get("title", ""),
            "company": company_name or "Our Company"
        }
        
        notification_results = []
        
        for result in self.match_results:
            candidate_info = next(
                (c for c in self.candidates if c["contact_info"]["email"] == result["email"]), 
                None
            )
            
            if candidate_info and candidate_info["contact_info"]["email"]:
                success = self.notification_system.notify_candidate(
                    candidate_info,
                    job_info,
                    result["details"]
                )
                
                notification_results.append({
                    "candidate": result["candidate"],
                    "email": result["email"],
                    "notification_sent": success,
                    "shortlisted": result["shortlisted"]
                })
        
        return notification_results
    
    def save_results(self, output_file):
        """Save screening results to JSON file"""
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(self.match_results, f, indent=2, ensure_ascii=False)