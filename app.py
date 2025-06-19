from flask import Flask, request, jsonify, render_template, redirect, url_for, flash, send_from_directory, Response
from werkzeug.utils import secure_filename
import os
import json
import traceback
from datetime import datetime

# Import the complete JobScreeningSystem and related classes
import re
# import numpy as np
# import pandas as pd
import spacy
import PyPDF2
import docx
import pdfplumber
import io
import csv
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# --- PDF Generation Imports ---
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
# --- End PDF Generation Imports ---

# Load NLP model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("Please install spacy English model: python -m spacy download en_core_web_sm")
    nlp = None

class JobDescriptionParser:
    """Extracts key requirements from job descriptions"""
    
    def __init__(self):
        pass
    
    def extract_text(self, file_path):
        """Extract text from JD files in various formats"""
        if file_path.endswith('.pdf'):
            return self._extract_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            return self._extract_from_docx(file_path)
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        else:
            raise ValueError("Unsupported file format")
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            app.logger.error(f"Error extracting from PDF with pdfplumber: {e}")
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e2:
                app.logger.error(f"Fallback PDF extraction with PyPDF2 also failed: {e2}")
        return text
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            app.logger.error(f"Error extracting from DOCX: {e}")
            return ""
    
    def summarize(self, text):
        """Generate simple summary of job description"""
        sentences = text.split('.')
        summary = '. '.join(sentences[:3]) + '.'
        return summary
    
    def extract_requirements(self, text):
        """Extract key requirements from job description"""
        if not nlp:
            return self._extract_requirements_basic(text)
        
        doc = nlp(text)
        
        skills = []
        education = []
        experience = []
        
        skill_patterns = ["skills", "proficiency", "knowledge of", "familiarity with", "expertise in"]
        edu_patterns = ["degree", "qualification", "education", "graduate", "bachelor", "master", "phd"]
        exp_patterns = ["experience", "years", "worked", "background in"]
        
        for sent in doc.sents:
            sent_text = sent.text.lower()
            
            if any(pattern in sent_text for pattern in skill_patterns):
                skills.append(sent.text.strip().capitalize())
            
            if any(pattern in sent_text for pattern in edu_patterns):
                education.append(sent.text.strip().capitalize())
            
            if any(pattern in sent_text for pattern in exp_patterns):
                experience.append(sent.text.strip().capitalize())
        
        location_pattern = r"location:?\s*([\w\s,]+)"
        job_type_pattern = r"(full[- ]time|part[- ]time|contract|remote|hybrid)"
        
        location_match = re.search(location_pattern, text, re.IGNORECASE)
        job_type_match = re.search(job_type_pattern, text, re.IGNORECASE)
        
        location = location_match.group(1).strip().capitalize() if location_match else "Not specified"
        job_type = job_type_match.group(1).strip().capitalize() if job_type_match else "Not specified"
        
        requirements = {
            "skills": skills,
            "education": education,
            "experience": experience,
            "location": location.capitalize(), # Ensure "Not specified" is also capitalized
            "job_type": job_type,
            "summary": self.summarize(text)
        }
        
        return requirements
    
    def _extract_requirements_basic(self, text):
        """Basic requirements extraction without spaCy"""
        text_lower = text.lower()
        
        skills = []
        skill_keywords = ["python", "java", "javascript", "sql", "html", "css", "react", "angular", "node"]
        for skill in skill_keywords:
            if skill.lower() in text_lower: # Ensure comparison is robust
                skills.append(skill.capitalize())
        
        education = []
        if "bachelor" in text_lower or "degree" in text_lower:
            education.append("Bachelor's degree required".capitalize())
        if "master" in text_lower:
            education.append("Master's degree preferred".capitalize())
        
        experience = []
        exp_match = re.search(r"(\d+)\s*years?\s*experience", text_lower)
        if exp_match:
            experience.append(f"{exp_match.group(1)} years of experience required".capitalize())
        
        return {
            "skills": skills,
            "education": education,
            "experience": experience,
            "location": "Not specified".capitalize(),
            "job_type": "Not specified".capitalize(),
            "summary": self.summarize(text)
        }


class ResumeParser:
    """Extracts structured information from candidate resumes"""
    
    def __init__(self):
        self.nlp = nlp
    
    def extract_text(self, file_path):
        """Extract text from resume in various formats"""
        if file_path.endswith('.pdf'):
            return self._extract_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            return self._extract_from_docx(file_path)
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        else:
            raise ValueError("Unsupported file format")
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            app.logger.error(f"Error extracting resume from PDF with pdfplumber: {e}")
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e2:
                app.logger.error(f"Fallback resume PDF extraction with PyPDF2 also failed: {e2}")
        return text
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            app.logger.error(f"Error extracting resume from DOCX: {e}")
            return ""
    
    def extract_contact_info(self, text):
        """Extract contact information from resume"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)

        phone_pattern = r'(\+\d{1,3}\s?)?(\(?\d{3}\)?[\s.-]?)?\d{3}[\s.-]?\d{4}'
        raw_phones = re.findall(phone_pattern, text)
        # Extract only digits and filter out very short sequences
        phones_digits = [''.join(filter(str.isdigit, ''.join(parts))) for parts in raw_phones]
        phones_digits = [p for p in phones_digits if len(p) >= 7] # Basic filter for length

        name = ""
        stripped_lines = [line.strip() for line in text.strip().split('\n') if line.strip()]

        # Attempt 1: Heuristic for name in the first few lines
        if not name and stripped_lines:
            exclusion_keywords = [
                'email', 'phone', 'mobile', 'tel', 'linkedin', 'github', 'portfolio', 'website', 'http', '@',
                'profile', 'summary', 'objective', 'career', 'experience', 'education', 'skills', 
                'projects', 'awards', 'references', 'contact', 'information', 'details', 'address',
                'january', 'february', 'march', 'april', 'may', 'june', 'july', 'august', 'september', 'october', 'november', 'december'
            ]
            for i in range(min(5, len(stripped_lines))): # Check first 5 non-empty lines
                line_text = stripped_lines[i]
                words = line_text.split()

                if not (1 <= len(words) <= 6 and len(line_text) < 60 and '@' not in line_text and 'http' not in line_text):
                    continue

                contains_exclusion = False
                for kw in exclusion_keywords:
                    if re.search(r'\b' + re.escape(kw) + r'\b', line_text, re.IGNORECASE):
                        contains_exclusion = True
                        break
                if contains_exclusion:
                    continue
                
                if re.fullmatch(r"[A-Za-zÀ-ÖØ-öø-ÿ\s'-]+", line_text) and not line_text.endswith(('.', '!', '?')):
                    is_capitalized_properly = False
                    if all(word[0].isupper() for word in words if word and word[0].isalpha()):
                        is_capitalized_properly = True
                    elif line_text.isupper() and len(words) <= 3:
                        is_capitalized_properly = True
                    
                    if is_capitalized_properly:
                        if len(words) > 1 or (len(words) == 1 and len(words[0]) > 2 and words[0].lower() not in ['resume', 'cv']):
                            name = line_text
                            app.logger.info(f"Name extracted by heuristic (top lines): '{name}'")
                            break 
        
        # Attempt 2: spaCy PERSON entity
        if not name and self.nlp:
            doc = self.nlp(text[:1000]) # Process only the beginning for relevance and speed
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    potential_name = ent.text.strip()
                    p_words = potential_name.split()
                    if 1 <= len(p_words) <= 5 and len(potential_name) < 50 and '.' not in potential_name and '@' not in potential_name:
                        if all(w[0].isupper() for w in p_words if w and w[0].isalpha()) or (len(p_words) <=2 and potential_name.isupper()):
                            name = potential_name
                            app.logger.info(f"Name extracted by spaCy: '{name}'")
                            break 
        
        # Attempt 3: Fallback to the first non-empty line (very constrained)
        if not name and stripped_lines:
            first_line_candidate = stripped_lines[0]
            fallback_exclusion_keywords = exclusion_keywords + ['resume', 'cv', 'curriculum vitae']
            if 1 <= len(first_line_candidate.split()) <= 5 and \
               len(first_line_candidate) < 50 and \
               not any(char.isdigit() for char in first_line_candidate) and \
               not any(kw.lower() in first_line_candidate.lower() for kw in fallback_exclusion_keywords) and \
               re.fullmatch(r"[A-Za-zÀ-ÖØ-öø-ÿ\s'-]+", first_line_candidate):
                name = first_line_candidate
                app.logger.info(f"Name extracted by fallback (first line): '{name}'")

        return {
            "name": name.strip(),
            "email": emails[0] if emails else "",
            "phone": phones_digits[0] if phones_digits else "" # Return the first, longest, cleaned phone number
        }
    
    def extract_education(self, text):
        """Extract education information from resume"""
        education = []
        edu_keywords = ["degree", "bachelor", "master", "phd", "btech", "mtech", "b.tech", "m.tech", "diploma"]
        
        if self.nlp:
            doc = self.nlp(text)
            for sent in doc.sents:
                sent_text = sent.text.lower()
                if any(keyword in sent_text for keyword in edu_keywords):
                    education.append(sent.text.strip().capitalize())
        else:
            lines = text.split('\n')
            for line in lines:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in edu_keywords):
                    education.append(line.strip().capitalize())
        
        return education
    
    def extract_skills(self, text):
        """Extract skills from resume"""
        tech_skills = [
            "python", "java", "javascript", "sql", "html", "css", "react", "angular", "node", 
            "aws", "azure", "gcp", "docker", "kubernetes", "git", "tensorflow", "pytorch",
            "machine learning", "data science", "deep learning", "ai", "artificial intelligence",
            "nlp", "natural language processing", "computer vision", "c++", "c#", "php", "r"
        ] # These are kept lowercase for matching
        
        skills = []
        text_lower = text.lower()
        
        for skill in tech_skills:
            # Use regex for whole word matching for better accuracy
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                skills.append(skill.capitalize()) # Capitalize the skill before adding
                
        return skills
    
    def extract_experience(self, text):
        """Extract work experience from resume"""
        experience = []
        exp_keywords = ["experience", "work", "employment", "job", "position", "role"]
        
        if self.nlp:
            doc = self.nlp(text)
            
            exp_section = None
            lines = text.split('\n')
            for i, line in enumerate(lines):
                if any(keyword in line.lower() for keyword in exp_keywords) and len(line) < 30:
                    exp_section = i
                    break
            
            if exp_section is not None:
                exp_text = '\n'.join(lines[exp_section+1:])
                for sent in self.nlp(exp_text).sents:
                    for ent in sent.ents:
                        if ent.label_ == "ORG":
                            experience.append(sent.text.strip().capitalize())
                            break
        else:
            lines = text.split('\n')
            for line in lines:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in exp_keywords):
                    experience.append(line.strip().capitalize())
        
        return experience
    
    def parse_resume(self, file_path):
        """Parse a resume file and extract structured information"""
        text = self.extract_text(file_path)
        
        contact_info = self.extract_contact_info(text)
        education = self.extract_education(text)
        skills = self.extract_skills(text)
        experience = self.extract_experience(text)
        
        return {
            "contact_info": contact_info,
            "education": education,
            "skills": skills,
            "experience": experience,
            "full_text": text
        }


class MatchingEngine:
    """Implements algorithms for matching resumes against job requirements"""
    
    def __init__(self):
        self.vectorizer = TfidfVectorizer(ngram_range=(1, 2))
    
    def calculate_skill_match(self, jd_skill_phrases, resume_canonical_skills):
        """Calculate match score for skills"""
        # If JD parsing yields no skill sentences, assume perfect match for this category
        if not jd_skill_phrases:
            return 1.0
        # If resume has no skills, but JD expects some (jd_skill_phrases is not empty)
        if not resume_canonical_skills:
            return 0.0

        jd_text_blob_lower = " ".join(s.lower() for s in jd_skill_phrases)
        
        matched_count = 0
        for r_skill in resume_canonical_skills: # e.g., "Python", "Machine Learning"
            # Match the canonical resume skill (lowercase, with word boundaries) in the JD text blob
            pattern = r'\b' + re.escape(r_skill.lower()) + r'\b'
            if re.search(pattern, jd_text_blob_lower):
                matched_count += 1
        
        # Score: proportion of the candidate's listed skills that are found in the JD's skill requirements.
        return matched_count / len(resume_canonical_skills)
    
    def calculate_experience_match(self, jd_experience, resume_experience):
        """Calculate match score for experience using semantic similarity"""
        if not jd_experience or not resume_experience:
            return 0.0
            
        jd_exp_text = ' '.join(jd_experience)
        resume_exp_text = ' '.join(resume_experience)
        
        try:
            corpus = [jd_exp_text, resume_exp_text]
            tfidf_matrix = self.vectorizer.fit_transform(corpus)
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            return similarity
        except Exception as e:
            app.logger.error(f"Error calculating experience match: {e}")
            return 0.0
    
    def calculate_education_match(self, jd_education, resume_education):
        """Calculate match score for education"""
        if not jd_education:
            return 1.0
            
        if not resume_education:
            return 0.0
        
        try:
            jd_edu_text = ' '.join(jd_education)
            resume_edu_text = ' '.join(resume_education)
            
            corpus = [jd_edu_text, resume_edu_text]
            tfidf_matrix = self.vectorizer.fit_transform(corpus)
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            return similarity
        except Exception as e:
            app.logger.error(f"Error calculating education match: {e}")
            return 0.0
    
    def match_resume_to_jd(self, job_requirements, resume_data, weights=None):
        """Calculate overall match score between resume and job description"""
        if weights is None:
            weights = {
                "skills": 0.5,
                "experience": 0.3,
                "education": 0.2
            }
        
        skill_score = self.calculate_skill_match(
            job_requirements.get("skills", []),
            resume_data.get("skills", [])
        )
        
        experience_score = self.calculate_experience_match(
            job_requirements.get("experience", []),
            resume_data.get("experience", [])
        )
        
        education_score = self.calculate_education_match(
            job_requirements.get("education", []),
            resume_data.get("education", [])
        )
        
        overall_score = (
            skill_score * weights["skills"] +
            experience_score * weights["experience"] +
            education_score * weights["education"]
        )
        
        match_details = {
            "overall_score": overall_score,
            "skill_score": skill_score,
            "experience_score": experience_score,
            "education_score": education_score,
            "matched_skills": []
        }
        
        # New logic for populating match_details["matched_skills"] for clarity
        jd_skill_phrases = job_requirements.get("skills", [])  # List of skill-related sentences from JD
        resume_skill_keywords = resume_data.get("skills", []) # List of canonical skills from resume, e.g., ["python", "machine learning"]

        identified_matches = set()
        if jd_skill_phrases and resume_skill_keywords:
            # Combine all JD skill phrases into a single lowercase string for efficient searching.
            jd_skills_text_blob_lower = " ".join(s.lower() for s in jd_skill_phrases)

            for r_skill_keyword in resume_skill_keywords: # Iterate through canonical skills from resume
                # Use regex to match whole words/phrases, accounting for special characters in skill names.
                # For example, r_skill_keyword "c++" becomes pattern r'\bc\+\+\b'
                pattern = r'\b' + re.escape(r_skill_keyword.lower()) + r'\b'
                if re.search(pattern, jd_skills_text_blob_lower):
                    identified_matches.add(r_skill_keyword) # Add the canonical skill name

        match_details["matched_skills"] = sorted(list(identified_matches))
        
        return match_details


class NotificationSystem:
    """Handles automated communications with candidates"""

    def __init__(self, smtp_server, smtp_port, sender_email, sender_password):
        self.smtp_server = smtp_server
        self.smtp_port = smtp_port
        self.sender_email = sender_email
        self.sender_password = sender_password

        # Email templates
        self.templates = {
            "shortlisted": """
Subject: Update Regarding Your Application for the {job_title} at {company_name}

<p>Dear {candidate_name},</p>

<p>We are pleased to inform you that your application for the <strong>the {job_title}</strong> position at <strong>{company_name}</strong> has progressed to the next stage. Our initial review, supported by our AI-powered screening system, indicates a strong alignment between your profile and the requirements for this role.</p>

<p>Your profile received an overall match score of <strong>{match_score}%</strong>.</p>

<p>Our Human Resources department will be in contact with you in the near future to discuss the subsequent steps in the selection process, which may include an interview.</p>

<p>We appreciate your interest in {company_name} and look forward to potentially speaking with you further.</p>

<p>Sincerely,</p>
<p>The {company_name} Recruitment Team</p>
""",

            "rejected": """
Subject: Regarding Your Application for the {job_title} at {company_name}

<p>Dear {candidate_name},</p>

<p>Thank you for your interest in the <strong>the {job_title}</strong> position at <strong>{company_name}</strong> and for taking the time to apply.</p>

<p>We have received a significant number of applications for this role. After a thorough review of all candidates, including an initial assessment by our AI-powered screening system, we regret to inform you that we will not be moving forward with your application at this time. While your qualifications are commendable, other candidates were deemed to more closely match the specific requirements for this particular position.</p>

<p>This decision does not reflect on your overall capabilities, and we encourage you to visit our careers page regularly for future opportunities that may align with your skills and experience.</p>

<p>We wish you the best in your job search.</p>

<p>Sincerely,</p>
<p>The {company_name} Recruitment Team</p>
"""
        }

    def send_email(self, recipient_email, template_name, params):
        template = self.templates.get(template_name)
        if not template:
            raise ValueError(f"Template '{template_name}' not found")
        
        # Extract subject and body from the template
        # Assumes "Subject: " is the first line of the template content
        try:
            subject_line_raw, html_body_raw = template.split('\n\n', 1)
            subject_line = subject_line_raw.replace('Subject: ', '').strip().format(**params)
            html_body = html_body_raw.format(**params)
        except ValueError:
            app.logger.error(f"Email template '{template_name}' is not formatted correctly (missing Subject or body).")
            return False

        message = MIMEMultipart()
        message["From"] = self.sender_email
        message["To"] = recipient_email
        message["Subject"] = subject_line
        message.attach(MIMEText(html_body, "html")) # Send as HTML
        try:
            with smtplib.SMTP(self.smtp_server, self.smtp_port) as server:
                server.starttls()
                server.login(self.sender_email, self.sender_password)
                server.send_message(message)
            app.logger.info(f"Email successfully handed off to SMTP server for: {recipient_email}, Subject: {subject_line}")
            return True
        except smtplib.SMTPException as e:
            app.logger.error(f"SMTP error sending email to {recipient_email}: {str(e)}")
            return False
        except Exception as e: # Catch any other unexpected errors during sending
            app.logger.error(f"Unexpected error sending email to {recipient_email}: {str(e)}")
            return False

    def notify_candidate(self, candidate_info, job_info, match_details, threshold=0.8):
        match_score = match_details["overall_score"] * 100
        recipient_email = candidate_info["contact_info"]["email"]
        candidate_name = candidate_info["contact_info"]["name"] or "Candidate"

        params = {
            "candidate_name": candidate_name,
            "job_title": job_info.get("title", "the position"),
            "company_name": job_info.get("company", "our company"),
            "match_score": round(match_score, 1)
        }

        template_name = "shortlisted" if match_score >= (threshold * 100) else "rejected"
        return self.send_email(recipient_email, template_name, params)


class JobScreeningSystem:
    """Main system class that coordinates the job screening process"""

    def __init__(self, smtp_config=None):
        self.jd_parser = JobDescriptionParser()
        self.resume_parser = ResumeParser()
        self.matching_engine = MatchingEngine()

        if smtp_config and all(smtp_config.values()): # Ensure all SMTP config values are present
            try:
                self.notification_system = NotificationSystem(**smtp_config)
                app.logger.info("NotificationSystem initialized.")
            except Exception as e:
                app.logger.error(f"Failed to initialize NotificationSystem: {e}")
                self.notification_system = None
        else:
            self.notification_system = None
            app.logger.warning("NotificationSystem not configured due to missing SMTP details.")

        self.job_requirements = None
        self.candidates = []
        self.match_results = []
        self.last_screening_time = None
        
    def load_job_description(self, file_path):
        """Load and parse job description"""
        text = self.jd_parser.extract_text(file_path)
        self.job_requirements = self.jd_parser.extract_requirements(text)

        filename = os.path.basename(file_path)
        job_title = os.path.splitext(filename)[0]
        self.job_requirements["title"] = job_title

        return self.job_requirements

    def load_resumes(self, directory_path):
        """Load and parse all resumes in directory"""
        self.candidates = []
        
        if not os.path.exists(directory_path):
            app.logger.error(f"Resume directory {directory_path} does not exist")
            return []

        for filename in os.listdir(directory_path):
            if filename.endswith(('.pdf', '.docx', '.txt')):
                file_path = os.path.join(directory_path, filename)
                try:
                    resume_data = self.resume_parser.parse_resume(file_path)
                    resume_data["file_name"] = filename
                    self.candidates.append(resume_data)
                    app.logger.info(f"Successfully parsed resume: {filename}")
                except Exception as e:
                    app.logger.error(f"Error parsing resume {filename}: {e}")

        return self.candidates

    def screen_candidates(self, threshold=0.8, weights=None):
        """Screen all candidates against job requirements"""
        if not self.job_requirements:
            raise ValueError("Job requirements not loaded")

        if not self.candidates:
            raise ValueError("No candidates loaded")

        self.match_results = []

        for candidate in self.candidates:
            match_details = self.matching_engine.match_resume_to_jd(
                self.job_requirements,
                candidate,
                weights
            )

            result = {
                "candidate": candidate["contact_info"]["name"] or "Unknown",
                "email": candidate["contact_info"]["email"],
                "file_name": candidate["file_name"],
                "match_score": match_details["overall_score"],
                "shortlisted": match_details["overall_score"] >= threshold,
                "details": match_details
            }

            self.match_results.append(result)

        self.match_results.sort(key=lambda x: x["match_score"], reverse=True)

        # --- Automated notification for shortlisted candidates ---
        if self.notification_system:
            app.logger.info("Attempting automated notifications for shortlisted candidates...")
            job_info = {
                "title": self.job_requirements.get("title", "the relevant position"),
                "company": "Our Company" # Or make this configurable
            }
            notification_log_automated = []
            for result in self.match_results:
                if result.get("shortlisted"):
                    candidate_data = next((c for c in self.candidates if c.get("file_name") == result.get("file_name")), None)
                    if candidate_data and candidate_data.get("contact_info", {}).get("email"):
                        try:
                            # The threshold used here is the same one used for shortlisting.
                            # notify_candidate will correctly pick the "shortlisted" template.
                            success = self.notification_system.notify_candidate(candidate_data, job_info, result["details"], threshold)
                            status = "Sent (Automated - Shortlisted)" if success else "Failed (SMTP - Automated Shortlisted)"
                            notification_log_automated.append({
                                "candidate": result["candidate"], 
                                "email": candidate_data["contact_info"]["email"],
                                "status": status,
                                "shortlisted": True
                            })
                        except Exception as e_notify:
                            app.logger.error(f"Automated notification failed for {result['candidate']}: {str(e_notify)}")
                            notification_log_automated.append({"candidate": result["candidate"], "email": candidate_data["contact_info"]["email"], "status": f"Error: {str(e_notify)}", "shortlisted": True})
                    else:
                        notification_log_automated.append({"candidate": result["candidate"], "email": "N/A", "status": "Skipped (Automated - no email for shortlisted)", "shortlisted": True})
            app.logger.info(f"Automated Notification Log: {notification_log_automated}")
        # --- End automated notification ---
        
        automated_notifications_summary = {
            "attempted": len(notification_log_automated),
            "successful": len([log for log in notification_log_automated if "Sent" in log.get("status", "")]),
            "failed": len([log for log in notification_log_automated if "Failed" in log.get("status", "") or "Error" in log.get("status", "")])
        }

        self.last_screening_time = datetime.now().isoformat()
        return self.match_results, automated_notifications_summary
    
    def send_notifications(self, company_name=None, threshold=0.8, only_shortlisted=False):
        """
        Send notifications to candidates.
        If only_shortlisted is True, only sends to shortlisted candidates.
        Otherwise, sends to all (shortlisted get "shortlisted" email, others get "rejected").
        """
        if not self.notification_system:
            app.logger.warning("Notification system not configured. Cannot send emails.")
            raise ValueError("Notification system not configured. Please check SMTP settings.")

        if not self.match_results:
            app.logger.warning("No screening results available to send notifications.")
            raise ValueError("No screening results available to send notifications.")

        job_info = {
            "title": self.job_requirements.get("title", "the relevant position"),
            "company": company_name or "Our Company" # Default company name
        }

        manual_notification_log = []
        for result in self.match_results:
            if only_shortlisted and not result.get("shortlisted"):
                manual_notification_log.append({
                    "candidate": result.get("candidate", "Unknown"),
                    "email": result.get("email", "N/A"),
                    "status": "Skipped (not shortlisted for this manual notification run)",
                    "shortlisted": False
                })
                continue

            candidate_data = next((c for c in self.candidates if c.get("file_name") == result.get("file_name")), None)
            
            if not candidate_data or not candidate_data.get("contact_info", {}).get("email"):
                manual_notification_log.append({
                    "candidate": result.get("candidate", "Unknown"), 
                    "email": result.get("email", "N/A"),
                    "status": f"Skipped (no email for candidate)", 
                    "shortlisted": bool(result.get("shortlisted", False))
                })
                continue
            
            try:
                # notify_candidate uses the threshold to determine template (shortlisted/rejected)
                # The `result.get("shortlisted")` status is based on the screening_threshold.
                # The `threshold` param here is for the notification system's decision logic.
                email_sent_successfully = self.notification_system.notify_candidate(
                    candidate_data, 
                    job_info, 
                    result["details"], 
                    threshold # This threshold determines which template (shortlisted/rejected) is used by notify_candidate
                )
                
                # Determine status message based on whether the candidate was shortlisted by screening
                # and if the email was successfully sent.
                if result.get("shortlisted"):
                    email_status_message = "Sent (Shortlisted Template)"
                else: # Candidate was not shortlisted by screening, so rejection template was used
                    email_status_message = "Sent (Rejection Template)"

                if not email_sent_successfully:
                    email_status_message = "Failed (SMTP issue, check server logs)"

                manual_notification_log.append({
                    "candidate": result["candidate"], 
                    "email": candidate_data["contact_info"]["email"],
                    "status": email_status_message, 
                    "shortlisted": bool(result.get("shortlisted", False)) # Reflects screening shortlist status
                })
            except Exception as e_manual_notify:
                app.logger.error(f"Manual notification failed for {result['candidate']}: {str(e_manual_notify)}")
                manual_notification_log.append({
                    "candidate": result["candidate"], 
                    "email": candidate_data["contact_info"]["email"],
                    "status": f"Error: {str(e_manual_notify)}", 
                    "shortlisted": bool(result.get("shortlisted", False))
                })
        
        app.logger.info(f"Manual Notification Log: {manual_notification_log}")
        return manual_notification_log


# Flask Application
app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('FLASK_SECRET_KEY', 'your-default-secret-key-for-local-dev') # Load from env var
# Use /tmp for uploads on Vercel (ephemeral storage)
VERCEL_TMP_FOLDER = '/tmp'
app.config['UPLOAD_FOLDER'] = os.path.join(VERCEL_TMP_FOLDER, 'uploads_ai_job_screening')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Create upload directories
# These will be created on-demand within routes for serverless environments
# os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True) # Not here
# os.makedirs(os.path.join(app.config['UPLOAD_FOLDER'], 'job_descriptions'), exist_ok=True) # Not here
# os.makedirs(os.path.join(app.config['UPLOAD_FOLDER'], 'resumes'), exist_ok=True) # Not here

# Allowed file extensions
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Global screening system instance
# Placeholder SMTP configuration - REPLACE with your actual details or load from environment/config
# Ensure your SMTP server allows less secure app access or use an app password if using Gmail
SMTP_CONFIG = {
    "smtp_server": os.environ.get("SMTP_SERVER", "smtp.gmail.com"), # e.g., "smtp.gmail.com"
    "smtp_port": int(os.environ.get("SMTP_PORT", 587)), # e.g., 587 for TLS
    "sender_email": os.environ.get("SENDER_EMAIL", "omkarproject7@gmail.com"), # e.g., "your.email@gmail.com"
    "sender_password": os.environ.get("SENDER_PASSWORD", "tjxraiwtfugurliu") # e.g., "your_app_password"
}

screening_system = JobScreeningSystem(smtp_config=SMTP_CONFIG)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload_job_description', methods=['POST'])
def upload_job_description():
    try:
        if 'job_description' not in request.files:
            return jsonify({'error': 'No job description file provided'}), 400
        
        file = request.files['job_description']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{timestamp}_{filename}"
            
            jd_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'job_descriptions')
            os.makedirs(jd_dir, exist_ok=True) # Create on demand
            file_path = os.path.join(jd_dir, filename)
            file.save(file_path)
            
            # Parse job description using the complete system
            job_requirements = screening_system.load_job_description(file_path)
            
            return jsonify({
                'message': 'Job description uploaded and parsed successfully',
                'job_requirements': job_requirements,
                'filename': filename
            })
        
        return jsonify({'error': 'Invalid file type'}), 400
    
    except Exception as e:
        return jsonify({'error': f'Error processing job description: {str(e)}'}), 500

@app.route('/upload_resumes', methods=['POST'])
def upload_resumes():
    try:
        if 'resumes' not in request.files:
            return jsonify({'error': 'No resume files provided'}), 400
        
        files = request.files.getlist('resumes')
        if not files or all(file.filename == '' for file in files):
            return jsonify({'error': 'No files selected'}), 400
        
        uploaded_files = []
        resume_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'resumes')
        os.makedirs(resume_dir, exist_ok=True) # Create on demand

        # Clear previous resumes from /tmp (optional, depends on desired behavior for ephemeral storage)
        if os.path.exists(resume_dir):
            for old_file in os.listdir(resume_dir):
                try:
                    os.remove(os.path.join(resume_dir, old_file))
                except OSError: # File might have been removed by another concurrent function or non-existent
                    pass
        # else: # This case is covered by the makedirs above
            # os.makedirs(resume_dir, exist_ok=True)

        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{timestamp}_{filename}"
                
                file_path = os.path.join(resume_dir, filename)
                file.save(file_path)
                uploaded_files.append(filename)
        
        if not uploaded_files:
            return jsonify({'error': 'No valid files uploaded'}), 400
        
        # Parse resumes using the complete system
        candidates = screening_system.load_resumes(resume_dir)
        
        return jsonify({
            'message': f'Successfully uploaded {len(uploaded_files)} resume(s)',
            'uploaded_files': uploaded_files,
            'candidates_parsed': len(candidates)
        })
    
    except Exception as e:
        return jsonify({'error': f'Error processing resumes: {str(e)}'}), 500

@app.route('/screen_candidates', methods=['POST'])
def screen_candidates():
    try:
        data = request.get_json() or {}
        threshold = float(data.get('threshold', 0.8))
        
        weights = data.get('weights', {
            'skills': 0.5,
            'experience': 0.3,
            'education': 0.2
        })
        
        if not screening_system.job_requirements:
            return jsonify({'error': 'No job description loaded. Please upload a job description first.'}), 400
        
        if not screening_system.candidates:
            return jsonify({'error': 'No candidates loaded. Please upload resumes first.'}), 400
        
        # Perform screening
        results, automated_notifications_summary = screening_system.screen_candidates(threshold, weights)
        
        # Convert results to JSON-serializable format
        serializable_results = []
        for result in results:
            serializable_result = {
                'candidate': str(result.get('candidate', 'Unknown')),
                'email': str(result.get('email', '')),
                'file_name': str(result.get('file_name', '')),
                'match_score': float(result.get('match_score', 0.0)),
                'shortlisted': bool(result.get('shortlisted', False)),
                'details': {
                    'overall_score': float(result['details'].get('overall_score', 0.0)),
                    'skill_score': float(result['details'].get('skill_score', 0.0)),
                    'experience_score': float(result['details'].get('experience_score', 0.0)),
                    'education_score': float(result['details'].get('education_score', 0.0)),
                    'matched_skills': [str(skill) for skill in result['details'].get('matched_skills', [])]
                }
            }
            serializable_results.append(serializable_result)
        
        # Clean job title for display (remove timestamp, extension, underscores, capitalize words)
        raw_title = screening_system.job_requirements.get('title', 'Job Title')
        clean_title = raw_title
        # Remove timestamp if present (e.g., 20230613_183053_Software_Developer)
        clean_title = re.sub(r'^\d{8}_\d{6}_', '', clean_title)
        # Remove extension if present
        clean_title = re.sub(r'\.[a-zA-Z0-9]+$', '', clean_title)
        # Replace underscores with spaces and capitalize words
        clean_title = ' '.join(word.capitalize() for word in clean_title.replace('_', ' ').split())
        
        response = {
            'message': 'Screening completed successfully',
            'total_candidates': int(len(serializable_results)),
            'shortlisted_candidates': int(len([r for r in serializable_results if r.get('shortlisted', False)])),
            'threshold': float(threshold),
            'results': serializable_results,
            'automated_notifications_summary': automated_notifications_summary,
            'job_title': clean_title
        }
        
        return jsonify(response)
    
    except Exception as e:
        traceback.print_exc()  # This will help debug the exact error
        return jsonify({'error': f'Error during screening: {str(e)}'}), 500

@app.route('/get_results')
def get_results():
    try:
        if not screening_system.match_results:
            return jsonify({'error': 'No screening results available'}), 400
        
        # Convert results to JSON-serializable format
        serializable_results = []
        for result in screening_system.match_results:
            serializable_result = {
                'candidate': str(result.get('candidate', 'Unknown')),
                'email': str(result.get('email', '')),
                'file_name': str(result.get('file_name', '')),
                'match_score': float(result.get('match_score', 0.0)),
                'shortlisted': bool(result.get('shortlisted', False)),
                'details': {
                    'overall_score': float(result['details'].get('overall_score', 0.0)),
                    'skill_score': float(result['details'].get('skill_score', 0.0)),
                    'experience_score': float(result['details'].get('experience_score', 0.0)),
                    'education_score': float(result['details'].get('education_score', 0.0)),
                    'matched_skills': [str(skill) for skill in result['details'].get('matched_skills', [])]
                }
            }
            serializable_results.append(serializable_result)
        
        return jsonify({
            'results': serializable_results,
            'total_candidates': int(len(serializable_results)),
            'shortlisted_candidates': int(len([r for r in serializable_results if r.get('shortlisted', False)]))
        })
    
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': f'Error retrieving results: {str(e)}'}), 500

@app.route('/send_notifications', methods=['POST'])
def send_notifications():
    try:
        data = request.get_json() or {}
        company_name = data.get('company_name', 'Our Company')
        threshold = float(data.get('threshold', 0.8)) # Get threshold from request or use default
        only_shortlisted_param = data.get('only_shortlisted', False) # New parameter
        if not screening_system.notification_system:
            return jsonify({'error': 'Notification system not configured. Cannot send emails.'}), 400

        if not screening_system.match_results:
            return jsonify({'error': 'No screening results available'}), 400

        # Call the actual send_notifications method
        # The method now returns a log of notification attempts
        notification_log = screening_system.send_notifications(
            company_name=company_name, 
            threshold=threshold,
            only_shortlisted=bool(only_shortlisted_param)
        )

        return jsonify({
            'message': 'Notification process completed. Check log for details.',
            'notification_log': notification_log
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': f'Error sending notifications: {str(e)}'}), 500

@app.route('/download_results')
def download_results():
    try:
        if not screening_system.match_results:
            flash('No results available to download.', 'warning')
            return redirect(url_for('index')) # Or return jsonify error

        # Ensure the base upload folder exists for writing the PDF
        base_results_dir = app.config['UPLOAD_FOLDER']
        os.makedirs(base_results_dir, exist_ok=True)

        job_title_raw = screening_system.job_requirements.get('title', 'Screening')
        # Basic sanitization for filename
        job_title_safe = "".join(c if c.isalnum() or c in [' ', '_', '-'] else '' for c in job_title_raw).strip().replace(' ', '_')
        if not job_title_safe: job_title_safe = "Screening_Results" # Fallback
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        dynamic_pdf_filename = f"{job_title_safe}_results_{timestamp}.pdf"
        pdf_file_path = os.path.join(base_results_dir, dynamic_pdf_filename)
        
        # Create styles
        styles = getSampleStyleSheet()
        title_style = styles['h1']
        title_style.alignment = 1 # Center alignment for title
        heading_style = styles['h2']
        heading_style.alignment = 1 # Center alignment for job title heading

        header_style = ParagraphStyle(
            'HeaderStyle', parent=styles['Normal'], fontName='Helvetica-Bold', 
            fontSize=9, leading=11, alignment=1 # Center alignment for headers
        )

        # Base font for the document
        base_font_name = 'Helvetica'
        base_bold_font_name = 'Helvetica-Bold'

        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontName=base_font_name,
            fontSize=9,
            leading=11,
            alignment=0 # Left alignment
        )
        small_normal_style = ParagraphStyle(
            'SmallNormal',
            parent=normal_style,
            fontName=base_font_name,
            fontSize=8,
            leading=10,
            alignment=0 # Left alignment
        )


        doc = SimpleDocTemplate(pdf_file_path, pagesize=landscape(letter),
                                rightMargin=0.5*inch, leftMargin=0.5*inch,
                                topMargin=0.5*inch, bottomMargin=0.5*inch)
        story = []

        story.append(Paragraph("Job Screening Results", title_style))
        story.append(Spacer(1, 0.2*inch))

        if screening_system.job_requirements and screening_system.job_requirements.get('title'):
            story.append(Paragraph(f"Job Title: {screening_system.job_requirements['title']}", heading_style))
            story.append(Spacer(1, 0.1*inch))

        # Table data
        table_data = [
            [Paragraph("Candidate", header_style), Paragraph("Email", header_style), 
             Paragraph("File Name", header_style), Paragraph("Match Score (%)", header_style), 
             Paragraph("Shortlisted?", header_style), Paragraph("Skill Score", header_style), 
             Paragraph("Exp. Score", header_style), Paragraph("Edu. Score", header_style), 
             Paragraph("Matched Skills", header_style)
            ]
        ]

        for result in screening_system.match_results:
            candidate_name = str(result.get('candidate', 'Unknown'))
            email = str(result.get('email', 'N/A'))
            file_name = str(result.get('file_name', 'N/A'))
            match_score = f"{float(result.get('match_score', 0.0) * 100):.1f}"
            shortlisted_text = "Yes" if bool(result.get('shortlisted', False)) else "No"
            
            details = result.get('details', {})
            skill_score_val = f"{float(details.get('skill_score', 0.0) * 100):.1f}"
            exp_score_val = f"{float(details.get('experience_score', 0.0) * 100):.1f}"
            edu_score_val = f"{float(details.get('education_score', 0.0) * 100):.1f}"
            
            matched_skills_list = details.get('matched_skills', [])
            matched_skills_str = ', '.join(map(str, matched_skills_list))
            
            # Use Paragraph for all cell content to ensure consistent styling and wrapping
            candidate_para = Paragraph(candidate_name, normal_style)
            email_para = Paragraph(email, small_normal_style)
            filename_para = Paragraph(file_name, small_normal_style)
            matched_skills_para = Paragraph(matched_skills_str, small_normal_style)

            # Center align scores and boolean-like text
            centered_normal_style = ParagraphStyle('CenteredNormal', parent=normal_style, alignment=1)
            match_score_para = Paragraph(match_score, centered_normal_style)
            shortlisted_para = Paragraph(shortlisted_text, centered_normal_style)
            skill_score_para = Paragraph(skill_score_val, centered_normal_style)
            exp_score_para = Paragraph(exp_score_val, centered_normal_style)
            edu_score_para = Paragraph(edu_score_val, centered_normal_style)

            table_data.append([
                candidate_para,
                email_para,
                filename_para,
                match_score_para,
                shortlisted_para,
                skill_score_para,
                exp_score_para,
                edu_score_para,
                matched_skills_para
            ])

        # Create table
        # Adjusted column widths to fit within landscape letter (10 inches usable width with 0.5in margins)
        results_table = Table(table_data, colWidths=[
            1.4*inch, 1.6*inch, 1.2*inch, # Candidate, Email, File Name
            0.8*inch, 0.7*inch,           # Match Score, Shortlisted
            0.8*inch, 0.8*inch, 0.8*inch, # Skill, Exp, Edu Scores
            1.9*inch                      # Matched Skills
        ]) # Total width approx 10.0 inches

        results_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'), # Header text centered
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), # Vertical alignment
            ('FONTNAME', (0, 0), (-1, 0), base_bold_font_name),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0,0), (-1,0), 6), # Padding for header
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('LEFTPADDING', (0,0), (-1,-1), 6), # Padding for all cells
            ('RIGHTPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,1), (-1,-1), 4), # Top padding for data cells
            ('BOTTOMPADDING', (0,1), (-1,-1), 4) # Bottom padding for data cells
        ]))
        story.append(results_table)
        
        doc.build(story)

        return send_from_directory(directory=base_results_dir, path=dynamic_pdf_filename, as_attachment=True, download_name=dynamic_pdf_filename)

    except Exception as e:
        traceback.print_exc()
        # flash(f'Error generating PDF: {str(e)}', 'danger')
        # return redirect(url_for('index'))
        return jsonify({'error': f'Error generating PDF report: {str(e)}'}), 500

@app.route('/download_results_csv')
def download_results_csv():
    try:
        if not screening_system.match_results:
            return jsonify({'error': 'No results available to download as CSV.'}), 400

        job_title_raw = screening_system.job_requirements.get('title', 'Screening')
        job_title_safe = "".join(c if c.isalnum() or c in [' ', '_', '-'] else '' for c in job_title_raw).strip().replace(' ', '_')
        if not job_title_safe: job_title_safe = "Screening_Results"
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        csv_filename = f"{job_title_safe}_results_{timestamp}.csv"

        output = io.StringIO()
        writer = csv.writer(output)

        header = ["Candidate", "Email", "File Name", "Match Score (%)", "Shortlisted?", 
                  "Skill Score (%)", "Experience Score (%)", "Education Score (%)", "Matched Skills"]
        writer.writerow(header)

        for result in screening_system.match_results:
            details = result.get('details', {})
            row = [
                str(result.get('candidate', 'Unknown')),
                str(result.get('email', 'N/A')),
                str(result.get('file_name', 'N/A')),
                f"{float(result.get('match_score', 0.0) * 100):.1f}",
                "Yes" if bool(result.get('shortlisted', False)) else "No",
                f"{float(details.get('skill_score', 0.0) * 100):.1f}",
                f"{float(details.get('experience_score', 0.0) * 100):.1f}",
                f"{float(details.get('education_score', 0.0) * 100):.1f}",
                ', '.join(map(str, details.get('matched_skills', [])))
            ]
            writer.writerow(row)

        output.seek(0)
        
        return Response(
            output,
            mimetype="text/csv",
            headers={"Content-Disposition": f"attachment;filename={csv_filename}"}
        )

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': f'Error generating CSV report: {str(e)}'}), 500

@app.route('/reset_system', methods=['POST'])
def reset_system():
    try:
        global screening_system
        # Re-initialize with SMTP_CONFIG
        screening_system = JobScreeningSystem(smtp_config=SMTP_CONFIG)

        # Clear /tmp subdirectories used by the app
        for folder_name in ['job_descriptions', 'resumes']:
            folder_path = os.path.join(app.config['UPLOAD_FOLDER'], folder_name)
            if os.path.exists(folder_path):
                for file_item in os.listdir(folder_path):
                    item_path = os.path.join(folder_path, file_item)
                    try:
                        if os.path.isfile(item_path) or os.path.islink(item_path):
                            os.unlink(item_path)
                        # If you expect subdirectories inside job_descriptions/resumes, add shutil.rmtree
                    except Exception as e_remove: # pragma: no cover
                        app.logger.warning(f"Error removing {item_path} during reset: {e_remove}")
        # Note: Dynamically named PDFs in app.config['UPLOAD_FOLDER'] are not specifically deleted here.
            
        return jsonify({'message': 'System reset successfully (ephemeral data in /tmp cleared)'})
    except Exception as e:
        return jsonify({'error': f'Error resetting system: {str(e)}'}), 500

@app.route('/status')
def get_status():
    try:
        status = {
            'job_description_loaded': screening_system.job_requirements is not None,
            'candidates_loaded': len(screening_system.candidates),
            'screening_completed': len(screening_system.match_results) > 0,
            'job_requirements': screening_system.job_requirements, # Could be large, consider summarizing
            'last_screening_time': screening_system.last_screening_time 
        }
        
        return jsonify(status)
    
    except Exception as e:
        return jsonify({'error': f'Error getting status: {str(e)}'}), 500

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File too large. Maximum size is 16MB.'}), 413

@app.errorhandler(404)
def not_found(e):
    return jsonify({'error': 'Endpoint not found'}), 404

@app.errorhandler(500)
def internal_error(e):
    # Log the error for server-side inspection
    traceback.print_exc()
    return jsonify({'error': 'Internal server error. Please check server logs.'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
